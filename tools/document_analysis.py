"""
Document Analysis Tool for Jarvis
Read and analyze PDF and Word documents
"""
from langchain.tools import tool
import os
import logging
from pathlib import Path


@tool
def read_pdf_document(file_path: str, page_number: int = None) -> str:
    """
    Read content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        page_number: Specific page to read (optional, 1-indexed)
    
    Examples:
        - "Read this PDF: documents/report.pdf"
        - "What's on page 3 of manual.pdf?"
        - "Summarize contract.pdf"
    
    Returns:
        Text content from PDF
    """
    try:
        # Try importing PyPDF2
        try:
            import PyPDF2
        except ImportError:
            return "❌ PyPDF2 not installed. Run: pip install PyPDF2"
        
        # Expand user path and resolve
        file_path = os.path.expanduser(file_path)
        
        if not os.path.exists(file_path):
            return f"❌ File not found: {file_path}"
        
        if not file_path.lower().endswith('.pdf'):
            return f"❌ Not a PDF file: {file_path}"
        
        # Open and read PDF
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            # Read specific page
            if page_number is not None:
                if page_number < 1 or page_number > total_pages:
                    return f"❌ Invalid page number. PDF has {total_pages} pages."
                
                page = pdf_reader.pages[page_number - 1]
                text = page.extract_text()
                
                return f"📄 PDF: {Path(file_path).name}\nPage {page_number}/{total_pages}:\n\n{text}"
            
            # Read all pages
            else:
                texts = []
                for i, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        texts.append(f"--- Page {i} ---\n{page_text}")
                
                full_text = "\n\n".join(texts)
                
                # Limit output size
                if len(full_text) > 5000:
                    full_text = full_text[:5000] + "\n\n... [Text truncated, use page_number parameter for specific pages]"
                
                return f"📄 PDF: {Path(file_path).name}\nTotal Pages: {total_pages}\n\n{full_text}"
                
    except Exception as e:
        logging.error(f"PDF read error: {e}")
        return f"❌ Error reading PDF: {str(e)}"


@tool
def read_word_document(file_path: str) -> str:
    """
    Read content from a Word document (.docx).
    
    Args:
        file_path: Path to the Word document
    
    Examples:
        - "Read this Word doc: documents/notes.docx"
        - "What's in report.docx?"
        - "Summarize proposal.docx"
    
    Returns:
        Text content from Word document
    """
    try:
        # Try importing python-docx
        try:
            import docx
        except ImportError:
            return "❌ python-docx not installed. Run: pip install python-docx"
        
        # Expand user path and resolve
        file_path = os.path.expanduser(file_path)
        
        if not os.path.exists(file_path):
            return f"❌ File not found: {file_path}"
        
        if not file_path.lower().endswith('.docx'):
            return f"❌ Not a Word document (.docx): {file_path}"
        
        # Open and read Word document
        doc = docx.Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract tables
        tables_text = []
        for table in doc.tables:
            table_text = "\n[Table]"
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text += f"\n{row_text}"
            tables_text.append(table_text)
        
        # Combine content
        full_text = "\n\n".join(paragraphs)
        
        if tables_text:
            full_text += "\n\n" + "\n\n".join(tables_text)
        
        # Limit output size
        if len(full_text) > 5000:
            full_text = full_text[:5000] + "\n\n... [Text truncated]"
        
        return f"📄 Word Doc: {Path(file_path).name}\n\n{full_text}"
        
    except Exception as e:
        logging.error(f"Word read error: {e}")
        return f"❌ Error reading Word document: {str(e)}"


@tool
def read_text_document(file_path: str) -> str:
    """
    Read content from a plain text file (.txt, .md, .log, etc.).
    
    Args:
        file_path: Path to the text file
    
    Examples:
        - "Read notes.txt"
        - "What's in README.md?"
        - "Show me log.txt"
    
    Returns:
        Text content from file
    """
    try:
        # Expand user path and resolve
        file_path = os.path.expanduser(file_path)
        
        if not os.path.exists(file_path):
            return f"❌ File not found: {file_path}"
        
        # Read file with multiple encoding attempts
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                
                # Limit output size
                if len(content) > 5000:
                    content = content[:5000] + "\n\n... [Text truncated]"
                
                return f"📄 File: {Path(file_path).name}\n\n{content}"
                
            except UnicodeDecodeError:
                continue
        
        return f"❌ Could not decode file: {file_path}"
        
    except Exception as e:
        logging.error(f"Text read error: {e}")
        return f"❌ Error reading file: {str(e)}"


@tool
def analyze_document(file_path: str) -> str:
    """
    Analyze a document and provide metadata/summary.
    
    Args:
        file_path: Path to the document
    
    Examples:
        - "Analyze report.pdf"
        - "Get info about document.docx"
    
    Returns:
        Document metadata and basic analysis
    """
    try:
        file_path = os.path.expanduser(file_path)
        
        if not os.path.exists(file_path):
            return f"❌ File not found: {file_path}"
        
        # Get file info
        file_stat = os.stat(file_path)
        file_size = file_stat.st_size
        file_ext = Path(file_path).suffix.lower()
        
        # Format size
        if file_size < 1024:
            size_str = f"{file_size} bytes"
        elif file_size < 1024 * 1024:
            size_str = f"{file_size / 1024:.2f} KB"
        else:
            size_str = f"{file_size / (1024 * 1024):.2f} MB"
        
        # Basic analysis
        info = [
            f"📄 Document Analysis: {Path(file_path).name}",
            f"Type: {file_ext}",
            f"Size: {size_str}",
        ]
        
        # Type-specific analysis
        if file_ext == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    info.append(f"Pages: {len(pdf_reader.pages)}")
                    
                    # Get metadata
                    if pdf_reader.metadata:
                        if pdf_reader.metadata.title:
                            info.append(f"Title: {pdf_reader.metadata.title}")
                        if pdf_reader.metadata.author:
                            info.append(f"Author: {pdf_reader.metadata.author}")
            except:
                pass
        
        elif file_ext == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                info.append(f"Paragraphs: {len(doc.paragraphs)}")
                info.append(f"Tables: {len(doc.tables)}")
                
                # Word count
                word_count = sum(len(para.text.split()) for para in doc.paragraphs)
                info.append(f"Words: ~{word_count}")
            except:
                pass
        
        elif file_ext in ['.txt', '.md', '.log']:
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    lines = content.count('\n') + 1
                    words = len(content.split())
                    info.append(f"Lines: {lines}")
                    info.append(f"Words: {words}")
                    info.append(f"Characters: {len(content)}")
            except:
                pass
        
        return "\n".join(info)
        
    except Exception as e:
        logging.error(f"Document analysis error: {e}")
        return f"❌ Analysis error: {str(e)}"


# Quick test
if __name__ == "__main__":
    print("Document analysis tools created.")
    print("Requires: pip install PyPDF2 python-docx")
    print("\nSupported formats:")
    print("- PDF (.pdf)")
    print("- Word (.docx)")
    print("- Text (.txt, .md, .log)")
